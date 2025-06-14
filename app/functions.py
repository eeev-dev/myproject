from docx import Document

def create_doc(doc, interns, year, group, duration, head_teacher):

    def replace_placeholder_in_paragraphs(doc, placeholder, new_text):
        for p in doc.paragraphs:
            if placeholder in p.text:
                # Собираем полный текст абзаца
                full_text = ''.join(run.text for run in p.runs)
                if placeholder in full_text:
                    new_full_text = full_text.replace(placeholder, new_text)
                    # очищаем все run-ы
                    for run in p.runs:
                        run.text = ''
                    # вставляем новый текст в первый run
                    if p.runs:
                        p.runs[0].text = new_full_text
                    else:
                        p.add_run(new_full_text)


    replace_placeholder_in_paragraphs(doc, '{{YEAR}}', year)
    replace_placeholder_in_paragraphs(doc, '{{HEAD_TEACHER}}', head_teacher)
    replace_placeholder_in_paragraphs(doc, '{{GROUP}}', group)
    replace_placeholder_in_paragraphs(doc, '{{DURATION}}', duration)


    # Функция замены плейсхолдера {{TABLE}} таблицей
    def replace_table_placeholder(doc, placeholder, data):
        for paragraph in doc.paragraphs:
            if placeholder in paragraph.text:
                # Удаляем абзац с плейсхолдером
                p_element = paragraph._element
                parent_element = p_element.getparent()
                next_element = p_element.getnext()
                parent_element.remove(p_element)

                # Создаём таблицу через документ
                table = doc.add_table(rows=1, cols=len(data[0]))
                table.style = 'Table Grid'

                # Заполняем заголовки
                hdr_cells = table.rows[0].cells
                for j, val in enumerate(data[0]):
                    hdr_cells[j].text = val

                # Заполняем строки
                for row_data in data[1:]:
                    row_cells = table.add_row().cells
                    for k, val in enumerate(row_data):
                        row_cells[k].text = val

                # Перемещаем таблицу после удалённого параграфа
                tbl_element = table._element
                if next_element is not None:
                    parent_element.insert(parent_element.index(next_element), tbl_element)
                else:
                    parent_element.append(tbl_element)

                break

    table_data = [('№', 'Студенттердин ААА', 'Мекемелердин аты')]
    for idx, intern in enumerate(interns, start=1):
        place = intern.place
        if (intern.place == 'Свое место практики' and intern.title != None):
                place = intern.title 
        table_data.append((str(idx), intern.name, place))

    replace_table_placeholder(doc, '{{TABLE}}', table_data)


def find_teacher_and_subject(s):
    start = s.find(']') + 1
    first = s.find('[')
    end = s.find('[', first + 1)

    if 0 < start < end:
        return s[start:end].strip()
    else:
        return ''
    
    
def extract_subject(s):
    caps = [i for i, c in enumerate(s) if c.isupper()]
    if len(caps) >= 2:
        return s[caps[0]:caps[1]]
    else:
        return ''
    

def extract_teacher(s):
    parts = s.split()
    return ' '.join(parts[-2:])